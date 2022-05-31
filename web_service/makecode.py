'''
1. 按照input.txt裡面列出的資料，產生qrcode圖片，並存儲至qrcode_images資料夾
2. 將qrcode_images資料夾中的qrcode圖片匯入到word(d:/xampp/htdocs/library/qrcode.docx)檔中，並設定圖片寬度為使用者輸入的值
3. (ajax_upload2.php)將word檔下載到用戶端
'''

import os
import zipfile

import cv2
import numpy as np
import qrcode  # 匯入模組
from PIL import Image

import docx
from docx.enum.section import WD_ORIENTATION  # 處理文件的直向/橫向
from docx.shared import Cm


qr = qrcode.QRCode(
    version=1,
    error_correction=qrcode.constants.ERROR_CORRECT_L,
    box_size=30,
    border=0,
)


def generateQrcode(message):
    qr.clear()
    qr.add_data(message)
    qr.make(fit=True)
    return qr.make_image(fill_color="black", back_color="white")


def file2zip(zip_file_name: str, file_names: list):
    """ 將多個文件夾中文件壓縮存儲為zip

    :param zip_file_name:   /root/Document/test.zip
    :param file_names:      ['/root/user/doc/test.txt', ...]
    :return: 
    """
    # 讀取寫入方式 ZipFile requires mode 'r', 'w', 'x', or 'a'
    # 壓縮方式  ZIP_STORED： 存儲； ZIP_DEFLATED： 壓縮存儲
    with zipfile.ZipFile(zip_file_name, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        for fn in file_names:
            parent_path, name = os.path.split(fn)

            # zipfile 內置提供的將文件壓縮存儲在.zip文件中， arcname即zip文件中存入文件的名稱
            # 給予的歸檔名為 arcname (默認情況下將與 filename 一致，但是不帶驅動器盤符並會移除開頭的路徑分隔符)

            zf.write(fn, arcname=name)

            # 等價於以下兩行代碼
            # 切換目錄， 直接將文件寫入。不切換目錄，則會在壓縮文件中創建文件的整個路徑
            # os.chdir(parent_path)
            # zf.write(name)


def push_in_word(width):

    doc = docx.Document()  # 建立文件，如果已存在文件，請填入文件名，然後把下面的程式碼decommand

    dirPath = 'd:/Yolo_v4/qrcode_images'  # 這邊改成qrcode圖片的解壓縮檔路徑
    p = doc.add_paragraph()
    r = p.add_run()
    for f in os.listdir(dirPath):
        filename = os.path.join(dirPath, f)
        if os.path.isfile(filename):
            r.add_picture(filename, width=Cm(width))

    for sec in doc.sections:
        sec.orientation = WD_ORIENTATION.LANDSCAPE
        sec.left_margin = Cm(1.27)
        sec.right_margin = Cm(1.27)
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)

    doc.save('d:/xampp/htdocs/library/qrcode.docx')  # 存檔


if __name__ == '__main__':

    dirPath = "D:/Yolo_v4/qrcode_images"  # 清空名為'qrcode_images'的文件夹

    for f in os.listdir(dirPath):
        filename = os.path.join(dirPath, f)
        if os.path.isfile(filename):
            try:
                os.remove(filename)
            except OSError as e:
                print(e)
    if os.listdir(dirPath) == []:
        print("Folder is empty")

    file_names = []
    # 把你要生成的QR Code的資料寫在input.txt中
    with open("D:/Yolo_v4/web_service/input.txt", "r") as f:
        for content in f.readlines():
            content = content.strip('\n').split("/")
            content, size = content[0], int(content[1])
            # print(content, '{:.1f} cm'.format(size))

            # generate QR code
            qr_code = np.asarray(generateQrcode(content).convert("RGB").convert('L'))

            H, W = qr_code.shape[:2]

            module_sz = int(H / 21)

            canvas = np.full((H*2 + module_sz*7, module_sz * (11+6)), 255, dtype=np.uint8)

            # canvas[module_sz+module_sz//2:-module_sz-module_sz//2,
            #        module_sz+module_sz//2:-module_sz-module_sz//2] = 0
            
            canvas[module_sz*2:-module_sz*2, module_sz*2:-module_sz*2] = 255
            
            canvas[3*module_sz:24*module_sz, 3*module_sz:14 *
                   module_sz] = qr_code[:, :module_sz*11]
            
            canvas[25*module_sz:46*module_sz, 4*module_sz:14 *
                   module_sz] = np.flip(qr_code[:, module_sz*11:], axis=0)

            mw = np.full((module_sz, module_sz), 255, dtype=np.uint8)
            mb = np.full((module_sz, module_sz), 0, dtype=np.uint8)

            black = True

            # for y in range(21):
            #     if black:
            #         canvas[25*module_sz+y*module_sz:25*module_sz +
            #                (y+1)*module_sz, 3*module_sz:3*module_sz+module_sz] = mb
            #     else:
            #         canvas[25*module_sz+y*module_sz:25*module_sz +
            #                (y+1)*module_sz, 3*module_sz:3*module_sz+module_sz] = mw
            #     black = not black

            cv2.imwrite(dirPath+'/'+content+'.jpg', canvas)
            file_names.append(dirPath+'/'+content+'.jpg')

    push_in_word(size)

    # file2zip("d:/xampp/htdocs/library/qrcode.zip", file_names) # 打包成zip檔
