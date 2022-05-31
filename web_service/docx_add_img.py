'''
(實驗用)
可以練習看看如何用python操作docx(我也不太會)
python -m pip install docx (反正一切跟這個套件有關)
'''
import os

import docx
from docx.enum.section import WD_ORIENTATION  # 處理文件的直向/橫向
from docx.shared import Cm


def push_in_word(width):

    doc = docx.Document()  # 建立文件，如果已存在文件，請填入文件名，然後把下面的程式碼decommand

    dirPath = 'd:/Yolo_v4/qrcode_images'  # 這邊改成qrcode圖片的解壓縮檔路徑
    p = doc.add_paragraph()
    r = p.add_run()
    for f in os.listdir(dirPath):
        filename = os.path.join(dirPath, f)
        if os.path.isfile(filename):
            r.add_picture(filename, width=Cm(width))

    for sec in doc.sections:
        sec.orientation = WD_ORIENTATION.LANDSCAPE
        sec.left_margin = Cm(1.27)
        sec.right_margin = Cm(1.27)
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)

    doc.save('d:/xampp/htdocs/library/qr_code.docx')  # 存檔
